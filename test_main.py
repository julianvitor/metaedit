import io
from docx import Document
from fastapi.testclient import TestClient

from main import app

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _make_docx_bytes(paragraph_text: str) -> io.BytesIO:
    doc = Document()
    doc.add_paragraph(paragraph_text)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def test_metadata_flow_updates_core_properties():
    file_obj = _make_docx_bytes("Conteúdo de teste para o Julian.")

    with TestClient(app) as client:
        response = client.post(
            "/edit-metadata/",
            files={"file": ("test.docx", file_obj, DOCX_MIME)},
            data={
                "title": "Título Teste",
                "creator": "Julian Carreiro",
                "keyword": "curriculo, backend",
                "description": "Currículo atualizado",
                "category": "Pessoal",
            },
        )

    assert response.status_code == 200
    assert response.headers.get("content-type", "").startswith(DOCX_MIME)
    assert "attachment" in response.headers.get("content-disposition", "")
    assert "editado_test.docx" in response.headers.get("content-disposition", "")

    result_doc = Document(io.BytesIO(response.content))
    assert result_doc.core_properties.title == "Título Teste"
    assert result_doc.core_properties.author == "Julian Carreiro"
    assert result_doc.core_properties.keywords == "curriculo, backend"
    assert result_doc.core_properties.comments == "Currículo atualizado"
    assert result_doc.core_properties.category == "Pessoal"


def test_root_serves_index_html():
    with TestClient(app) as client:
        response = client.get("/")

    assert response.status_code == 200
    assert response.headers.get("content-type", "").startswith("text/html")